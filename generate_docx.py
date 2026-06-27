import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SCREENSHOTS_DIR = r"d:\Coding\projects\Internship\Spotter\screenshots"
OUTPUT_FILE = r"d:\Coding\projects\Internship\Spotter\Spotter_ELD_Submission.docx"

# ---------- Colour palette ----------
DARK_BG    = RGBColor(0x0F, 0x17, 0x2A)   # deep navy
ACCENT     = RGBColor(0x25, 0x63, 0xEB)   # electric blue
ACCENT2    = RGBColor(0x06, 0xB6, 0xD4)   # cyan
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF1, 0xF5, 0xF9)
MID_GRAY   = RGBColor(0x64, 0x74, 0x8B)
TEXT_DARK  = RGBColor(0x1E, 0x29, 0x3B)

# ---------- Helper: set paragraph shading ----------
def shade_paragraph(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def set_cell_bg(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell, top='none', bottom='none', left='none', right='none', color='auto', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:color'), color)
        borders.append(border)
    tcPr.append(borders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_run_font(run, name='Calibri', size=None, bold=False, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def set_paragraph_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_horizontal_rule(doc, color_hex='2563EB', thickness=12, before=60, after=60):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pb.append(bottom)
    pPr.append(pb)
    set_paragraph_spacing(p, before=before, after=after)
    return p

# ---------- Build the document ----------
doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.top_margin    = Inches(0.6)
section.bottom_margin = Inches(0.7)
section.left_margin   = Inches(0.85)
section.right_margin  = Inches(0.85)

# ============================
#  HERO HEADER (dark banner)
# ============================
hero_table = doc.add_table(rows=1, cols=1)
remove_table_borders(hero_table)
hero_cell = hero_table.rows[0].cells[0]
set_cell_bg(hero_cell, '0F172A')
hero_cell.width = Inches(7.8)

# Title line
hp = hero_cell.add_paragraph()
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(hp, before=140, after=20)
run = hp.add_run('SPOTTER ELD')
set_run_font(run, 'Calibri', size=38, bold=True, color=WHITE)

# Subtitle
sp = hero_cell.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(sp, before=0, after=10)
run2 = sp.add_run('Trip Planner & Log Generator')
set_run_font(run2, 'Calibri', size=14, color=RGBColor(0x93, 0xC5, 0xFD))

# Tag line
tp = hero_cell.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(tp, before=0, after=140)
run3 = tp.add_run('FMCSA HOS Compliant  •  Full Stack Django + React  •  Assessment Submission')
set_run_font(run3, 'Calibri', size=9, color=RGBColor(0x60, 0x7A, 0xA0), italic=True)

doc.add_paragraph()

# ============================
#  ACCENT STRIPE (blue bar)
# ============================
stripe = doc.add_table(rows=1, cols=3)
remove_table_borders(stripe)
widths  = [Inches(2.5), Inches(2.8), Inches(2.5)]
colors  = ['2563EB', '1E40AF', '06B6D4']
labels  = ['Django + React', 'FMCSA HOS Compliant', 'Full Stack App']

for i, (cell, w, c, lbl) in enumerate(zip(stripe.rows[0].cells, widths, colors, labels)):
    cell.width = w
    set_cell_bg(cell, c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=60, after=60)
    run = p.add_run(lbl)
    set_run_font(run, 'Calibri', size=10, bold=True, color=WHITE)

doc.add_paragraph()

# ============================
#  SECTION HELPER
# ============================
def section_heading(doc, title, emoji=''):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=120, after=30)
    if emoji:
        run0 = p.add_run(emoji + '  ')
        set_run_font(run0, 'Segoe UI Emoji', size=14)
    run = p.add_run(title.upper())
    set_run_font(run, 'Calibri', size=13, bold=True, color=ACCENT)
    add_horizontal_rule(doc, color_hex='2563EB', thickness=8, before=0, after=60)

# ============================
#  OVERVIEW
# ============================
section_heading(doc, 'Project Overview', '📋')
body = doc.add_paragraph(
    'Spotter is a comprehensive full-stack trip planner and log generator built for truck drivers. '
    'It accepts trip details, computes the optimal route via a free mapping API, and automatically '
    'produces FMCSA-compliant ELD Daily Log Sheets — eliminating manual calculation errors and '
    'ensuring legal compliance on every trip.'
)
set_paragraph_spacing(body, before=0, after=80)
for run in body.runs:
    set_run_font(run, 'Calibri', size=11, color=TEXT_DARK)

# ============================
#  DELIVERABLES (checklist table)
# ============================
section_heading(doc, 'Deliverables', '✅')

deliverables = [
    ('Live Hosted Version',   'Deployed via Vercel (frontend) + Railway/Render (backend)'),
    ('3-5 Min Loom Video',    'Walkthrough of the app and underlying code — script attached'),
    ('GitHub Repository',     'Full source code with README and setup instructions'),
]

dl_table = doc.add_table(rows=len(deliverables), cols=2)
remove_table_borders(dl_table)
for i, (label, detail) in enumerate(deliverables):
    row = dl_table.rows[i]
    row.cells[0].width = Inches(2.2)
    row.cells[1].width = Inches(5.6)
    bg = 'EFF6FF' if i % 2 == 0 else 'F8FAFC'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)

    # Label cell
    lp = row.cells[0].paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(lp, before=60, after=60)
    lr = lp.add_run('  ✔  ' + label)
    set_run_font(lr, 'Calibri', size=10, bold=True, color=ACCENT)

    # Detail cell
    dp = row.cells[1].paragraphs[0]
    set_paragraph_spacing(dp, before=60, after=60)
    dr = dp.add_run(detail)
    set_run_font(dr, 'Calibri', size=10, color=TEXT_DARK)

doc.add_paragraph()

# ============================
#  INPUTS & OUTPUTS (2-col)
# ============================
section_heading(doc, 'Inputs & Outputs', '🔄')

io_table = doc.add_table(rows=1, cols=2)
remove_table_borders(io_table)
left_cell  = io_table.rows[0].cells[0]
right_cell = io_table.rows[0].cells[1]
left_cell.width  = Inches(3.8)
right_cell.width = Inches(3.8)

# Inputs
set_cell_bg(left_cell, 'EFF6FF')
ih = left_cell.add_paragraph()
set_paragraph_spacing(ih, before=80, after=20)
irun = ih.add_run('  INPUTS')
set_run_font(irun, 'Calibri', size=10, bold=True, color=ACCENT)

inputs = ['Current Location', 'Pickup Location', 'Dropoff Location', 'Current Cycle Used (Hrs)']
for item in inputs:
    ip = left_cell.add_paragraph()
    set_paragraph_spacing(ip, before=20, after=20)
    ipr = ip.add_run(f'    →  {item}')
    set_run_font(ipr, 'Calibri', size=10, color=TEXT_DARK)

sp = left_cell.add_paragraph()
set_paragraph_spacing(sp, before=80, after=0)

# Outputs
set_cell_bg(right_cell, 'ECFEFF')
oh = right_cell.add_paragraph()
set_paragraph_spacing(oh, before=80, after=20)
orun = oh.add_run('  OUTPUTS')
set_run_font(orun, 'Calibri', size=10, bold=True, color=RGBColor(0x06, 0x82, 0x8D))

outputs = ['Interactive Map with route & stop info', 'Daily ELD Log Sheets (multi-day)', 'HOS Compliance breakdown', 'Fueling stop reminders']
for item in outputs:
    op = right_cell.add_paragraph()
    set_paragraph_spacing(op, before=20, after=20)
    opr = op.add_run(f'    →  {item}')
    set_run_font(opr, 'Calibri', size=10, color=TEXT_DARK)

sp2 = right_cell.add_paragraph()
set_paragraph_spacing(sp2, before=80, after=0)

doc.add_paragraph()

# ============================
#  TECH STACK
# ============================
section_heading(doc, 'Technology Stack', '⚙️')

tech = [
    ('Frontend',  'React.js · Vite · Axios · Leaflet.js'),
    ('Backend',   'Django · Django REST Framework · Python'),
    ('Mapping',   'OpenRouteService API (free tier)'),
    ('HOS Logic', 'Custom HOS Calculator — 70hr/8-day rules, 30-min break, 10-hr reset'),
    ('Hosting',   'Vercel (frontend) · Railway / Render (backend)'),
]

tech_table = doc.add_table(rows=len(tech), cols=2)
remove_table_borders(tech_table)
for i, (cat, detail) in enumerate(tech):
    row = tech_table.rows[i]
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(6.0)
    set_cell_bg(row.cells[0], '0F172A')
    set_cell_bg(row.cells[1], 'F8FAFC' if i % 2 == 0 else 'F1F5F9')

    lp = row.cells[0].paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(lp, before=60, after=60)
    lr = lp.add_run(cat + '  ')
    set_run_font(lr, 'Calibri', size=10, bold=True, color=ACCENT2)

    dp = row.cells[1].paragraphs[0]
    set_paragraph_spacing(dp, before=60, after=60)
    dr = dp.add_run('  ' + detail)
    set_run_font(dr, 'Calibri', size=10, color=TEXT_DARK)

doc.add_paragraph()

# ============================
#  ASSUMPTIONS
# ============================
section_heading(doc, 'Assumptions Applied', '📌')

assumptions = [
    'Property-carrying driver operating under the 70 hr / 8-day cycle.',
    'No adverse driving conditions modifier applied.',
    'Fueling stop inserted at least every 1,000 miles.',
    '1 hour allocated for pickup and 1 hour for drop-off.',
    '30-minute mandatory break after 8 continuous hours of driving.',
    '10-hour off-duty reset required before next driving window.',
]
for a in assumptions:
    ap = doc.add_paragraph()
    set_paragraph_spacing(ap, before=20, after=20)
    ar = ap.add_run('  •  ' + a)
    set_run_font(ar, 'Calibri', size=11, color=TEXT_DARK)

doc.add_paragraph()

# ============================
#  SCREENSHOTS
# ============================
section_heading(doc, 'Application Screenshots', '🖼️')

screenshots = [
    ("Main Interface — Empty State",       "Screenshot (2301).png",
     "The landing view with the input form on the left and a ready-to-plan prompt on the right."),
    ("Loading State",                       "Screenshot (2302).png",
     "Animated spinner while the backend computes routes and HOS schedule."),
    ("Route Map View",                     "Screenshot (2303).png",
     "Interactive Leaflet map showing the planned route with waypoints."),
    ("Route Map — Full Route",             "Screenshot (2304).png",
     "Zoomed-out map view displaying the complete multi-segment route."),
    ("ELD Log Sheet — Day 1",              "Screenshot (2305).png",
     "Auto-generated FMCSA Daily Log Sheet with graphed duty statuses for Day 1."),
    ("ELD Log Sheet — Day 2",              "Screenshot (2306).png",
     "Day 2 log sheet showing mandatory 30-minute break and drive segments."),
    ("Trip Summary & HOS Breakdown",       "Screenshot (2307).png",
     "Sidebar trip summary with distance, duration, and HOS compliance stats."),
]

for (title, filename, caption) in screenshots:
    img_path = os.path.join(SCREENSHOTS_DIR, filename)

    # screenshot label
    label_p = doc.add_paragraph()
    set_paragraph_spacing(label_p, before=80, after=10)
    lr = label_p.add_run(f'  {title}')
    set_run_font(lr, 'Calibri', size=11, bold=True, color=ACCENT)

    if os.path.exists(img_path):
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(img_p, before=0, after=6)
        run = img_p.add_run()
        run.add_picture(img_path, width=Inches(6.8))
    else:
        err_p = doc.add_paragraph(f'[Image not found: {filename}]')

    cap_p = doc.add_paragraph()
    set_paragraph_spacing(cap_p, before=4, after=60)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap_p.add_run(caption)
    set_run_font(cr, 'Calibri', size=9, italic=True, color=MID_GRAY)

# ============================
#  FOOTER BANNER
# ============================
doc.add_paragraph()
footer_table = doc.add_table(rows=1, cols=1)
remove_table_borders(footer_table)
fc = footer_table.rows[0].cells[0]
set_cell_bg(fc, '0F172A')

fp = fc.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(fp, before=80, after=80)
fr = fp.add_run('Spotter ELD  |  Full Stack Developer Assessment  |  Django + React  |  FMCSA Compliant')
set_run_font(fr, 'Calibri', size=9, color=RGBColor(0x60, 0x7A, 0xA0))

# ============================
#  SAVE
# ============================
doc.save(OUTPUT_FILE)
print(f'Saved: {OUTPUT_FILE}')
